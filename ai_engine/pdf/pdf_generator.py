"""
Resume generator using python-docx + docx2pdf.

Why this beats ReportLab:
  ReportLab:    Manual X/Y coordinate positioning. Hard to style. Brittle.
  python-docx:  Word handles ALL layout automatically. LibreOffice converts
                to pixel-perfect PDF in one CLI command. Zero layout code.

Pipeline:
  resume_text (str)
       ↓
  python-docx  →  resume.docx  (professional Word doc, ATS-safe)
       ↓
  docx2pdf    →  resume.pdf (perfect MS Word conversion)
       ↓
  FastAPI serves PDF bytes to frontend

Fallback:
  If MS Word is not installed → serve .docx directly.
  Most ATS systems accept .docx equally well as PDF.
"""

import os
import re
import shutil
import subprocess
import tempfile
import logging
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# ─── Design constants ─────────────────────────────────────────────────────────
FONT_NAME     = "Calibri"                   # universally supported by all ATS systems
NAME_SIZE     = Pt(18)
SECTION_SIZE  = Pt(11)
BODY_SIZE     = Pt(10)
CONTACT_SIZE  = Pt(9)
DATE_SIZE     = Pt(9)

HEADER_COLOR  = "1E3A5F"                    # navy hex for XML
SECTION_COLOR = RGBColor(0x1E, 0x3A, 0x5F)
BODY_COLOR    = RGBColor(0x2A, 0x2A, 0x2A)
WHITE_COLOR   = RGBColor(0xFF, 0xFF, 0xFF)
SUBTLE_COLOR  = RGBColor(0x55, 0x55, 0x55)

MARGIN_TOP    = Cm(1.2)
MARGIN_BOTTOM = Cm(1.2)
MARGIN_LEFT   = Cm(2.0)
MARGIN_RIGHT  = Cm(2.0)

SECTION_HEADERS = {
    "summary", "professional summary", "objective", "profile", "about",
    "experience", "work experience", "professional experience",
    "employment history", "work history", "career history",
    "education", "academic background", "qualifications",
    "skills", "technical skills", "core competencies", "technologies", "expertise",
    "projects", "personal projects", "academic projects",
    "certifications", "certificates", "licenses",
    "achievements", "awards", "honors", "accomplishments",
    "publications", "languages", "interests", "volunteer", "extra-curricular"
}

BULLET_RE   = re.compile(r'^([\s]*)([•\-\*●◦▪▸►✓✔◆■□▶→]+[ \t]+)(.*)')
NUMBERED_RE = re.compile(r'^\d+[.):]\s+')
DATE_RE     = re.compile(
    r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|Present|Current|20\d{2}|19\d{2})\b',
    re.IGNORECASE
)

_ACTION_VERBS = {
    'built','developed','designed','implemented','created','led','managed',
    'delivered','engineered','optimized','reduced','increased','automated',
    'architected','launched','collaborated','spearheaded','established',
    'maintained','improved','integrated','deployed','migrated','refactored',
    'analyzed','researched','trained','mentored','coordinated','worked',
    'helped','assisted','participated','contributed','performed','handled',
    'supported','drove','owned','shipped','scaled','streamlined',
}


# ─── Public API ───────────────────────────────────────────────────────────────

def generate_pdf(resume_text: str, output_path: str, candidate_name: str = "Candidate") -> None:
    """Generate professional PDF via docx + docx2pdf."""
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "resume.docx")
        _build_docx(resume_text, docx_path, candidate_name)
        pdf_path = _convert_to_pdf(docx_path, tmpdir)
        if pdf_path and os.path.exists(pdf_path):
            shutil.copy2(pdf_path, output_path)
        else:
            raise FileNotFoundError("PDF conversion failed. MS Word might not be installed.")


def generate_docx(resume_text: str, output_path: str, candidate_name: str = "Candidate") -> None:
    """Generate .docx only."""
    _build_docx(resume_text, output_path, candidate_name)


def is_pdf_conversion_available() -> bool:
    try:
        import win32com.client
        return True
    except ImportError:
        return False


# ─── PDF conversion ───────────────────────────────────────────────────

def _convert_to_pdf(docx_path: str, output_dir: str) -> str | None:
    try:
        from docx2pdf import convert
        pdf_path = os.path.join(output_dir, Path(docx_path).stem + ".pdf")
        # docx2pdf convert supports outputting directly to file
        convert(docx_path, pdf_path)
        return pdf_path if os.path.exists(pdf_path) else None
    except Exception as e:
        logger.error(f"Docx2Pdf conversion failed: {e}")
        return None


# ─── DOCX builder ─────────────────────────────────────────────────────────────

def _build_docx(resume_text: str, output_path: str, candidate_name: str) -> None:
    doc = Document()
    _set_page_margins(doc)
    _set_default_style(doc)

    raw_lines = resume_text.split('\n')
    lines = [_clean_line(l) for l in raw_lines]
    
    name_line = _find_name_line(lines)
    name_added = False
    i = 0

    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1; continue

        # ── Name / Header ──────────────────────────────────────────────────
        if not name_added and (line == name_line or i < 3):
            header_text = line if line else candidate_name
            _add_name_header(doc, header_text)
            name_added = True
            i += 1
            # Look for contact line
            if i < len(lines) and _is_contact_line(lines[i]):
                _add_contact_bar(doc, lines[i])
                i += 1
            continue

        # ── Section Heading ────────────────────────────────────────────────
        if _is_section_heading(line):
            _add_section_heading(doc, line)
            i += 1; continue

        # ── Job Header (Role | Company | Date) ─────────────────────────────
        if _is_job_header(line):
            _add_job_header(doc, line)
            i += 1; continue

        # ── Bullet Point ──────────────────────────────────────────────────
        bullet_match = BULLET_RE.match(line)
        if bullet_match or NUMBERED_RE.match(line):
            text = bullet_match.group(3) if bullet_match else NUMBERED_RE.sub('', line)
            _add_bullet(doc, text.strip())
            i += 1; continue

        # ── Standard Body Text ─────────────────────────────────────────────
        _add_body_text(doc, line)
        i += 1

    doc.save(output_path)


# ─── Paragraph builders ───────────────────────────────────────────────────────

def _clean_line(line: str) -> str:
    """Strip markdown bolding and hash headers."""
    l = line.strip()
    l = re.sub(r'\*\*|\*|__|_', '', l) # Strip markdown bold/italic
    l = l.lstrip('#').strip()           # Strip hash headers
    return l


def _set_page_margins(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin, section.bottom_margin = MARGIN_TOP, MARGIN_BOTTOM
        section.left_margin, section.right_margin = MARGIN_LEFT, MARGIN_RIGHT


def _set_default_style(doc: Document) -> None:
    style = doc.styles['Normal']
    style.font.name, style.font.size = FONT_NAME, BODY_SIZE
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(2)


def _add_name_header(doc: Document, name: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(4)
    _set_paragraph_shading(para, HEADER_COLOR)
    run = para.add_run(name.upper())
    run.font.size, run.font.bold = NAME_SIZE, True
    run.font.color.rgb = WHITE_COLOR


def _add_contact_bar(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(10)
    # Convert "Address | Email | Phone" into uniform look
    parts = [p.strip() for p in re.split(r'[|•\-\*]', text)]
    run = para.add_run("  •  ".join(filter(None, parts)))
    run.font.size, run.font.color.rgb = CONTACT_SIZE, SUBTLE_COLOR


def _add_section_heading(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after  = Pt(4)
    _add_paragraph_border(para)
    run = para.add_run(text.upper())
    run.font.size, run.font.bold = SECTION_SIZE, True
    run.font.color.rgb = SECTION_COLOR


def _add_job_header(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(1)
    
    # Try to split by common separators
    parts = [p.strip() for p in re.split(r'[|–\-\*•]', text)]
    
    if len(parts) >= 2:
        # Role in bold
        r1 = para.add_run(parts[0])
        r1.font.bold = True
        # Spacer
        para.add_run("  |  ").font.color.rgb = SUBTLE_COLOR
        # Company and Date in italic
        r2 = para.add_run(", ".join(parts[1:]))
        r2.font.italic = True
        r2.font.color.rgb = SUBTLE_COLOR
    else:
        para.add_run(text).font.bold = True


def _add_bullet(doc: Document, text: str) -> None:
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.left_indent = Inches(0.25)
    para.add_run(text)


def _add_body_text(doc: Document, text: str) -> None:
    para = doc.add_paragraph(text)
    para.paragraph_format.space_after = Pt(4)


# ─── XML & Classifier helpers ─────────────────────────────────────────────────

def _set_paragraph_shading(para, hex_color: str) -> None:
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def _add_paragraph_border(para) -> None:
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6'); bot.set(qn('w:space'), '1'); bot.set(qn('w:color'), HEADER_COLOR)
    pBdr.append(bot); pPr.append(pBdr)

def _find_name_line(lines: list[str]) -> str | None:
    for line in lines[:5]:
        s = line.strip()
        if not s: continue
        words = s.split()
        if 2 <= len(words) <= 5 and all(w[0].isupper() for w in words if w) and not any(c.isdigit() for c in s):
            return s
    return None

def _is_section_heading(line: str) -> bool:
    clean = line.lower().strip().rstrip(':')
    return clean in SECTION_HEADERS and len(line.split()) <= 4

def _is_contact_line(line: str) -> bool:
    return bool(re.search(r'[@]|linkedin|github|http|\+\d|\d{3}[-.\s]\d{4}', line, re.IGNORECASE))

def _is_job_header(line: str) -> bool:
    return (bool(DATE_RE.search(line)) or bool(re.search(r'[|–]', line))) and len(line.split()) <= 12
